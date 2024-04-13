from docx import Document
from os import path, listdir
import pandas as pd

def to_excel(pathto, filename):
    db = {}
    doc = Document(path.join(pathto, filename))
    db["IME DREVESA"] = doc.paragraphs[0].text
    db["LATINSKO IME"] = ""
    db["ANGLEŠKO IME"] = ""
    db["DRUŽINA"] = ""
    even = True

    lastheading = ""
    for para in doc.paragraphs[1:]:
        if("Heading" in para.style.name):
            lastheading = para.text
        else:
            if(lastheading != "" and para.text != "" and para.text != " "):
                try:
                    db[lastheading] += "\n"+para.text
                except:
                    db[lastheading] = para.text
    out = pd.DataFrame(data=db, index=[0])
    out.to_excel(path.join(pathto, (path.splitext(file)[0])+".xlsx"))
    #print(db["VIRI"])


if __name__ == "__main__":
    PATH = "docx/"
    for file in listdir(PATH):
        if file.endswith(".docx"):
            print(f"Pretvarjam datoteko {PATH}{file} v  {path.join(PATH, path.splitext(file)[0])}.xlsx")
            to_excel(PATH,file)